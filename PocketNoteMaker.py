"""

Warning:
This script was created for windows operating systems.
It doesn't work on Linux distributions.

Description:
This is a simple python script that takes a PDF file, 
creates images for each page in the pdf, stores those
in a folder and from that folder, merges 4 pictures 
in each docx page.

Dependencies: 
- docx
- pdf2image
- The path for poppler binaries (poppler_path) in line 69 
should be accurate; please make sure the path is alright. 
If needed, change it.

Set these variables before running this program:
- path
- pdf_file_name
- standard_file_path
- w_i
- h_i

"""



import os
import shutil
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = '' # Provide the path where the PDF file is, not the file itself. 
pdf_file_name = '' # Provide the name of the PDF file.

standard_file_path = f'' # Provide a blank docx file (full path). It shouldn't be in the folder where the PDF file is.

w_i = 3 # Provide the desired width for each small page inside the docx pages in INCHES
h_i = 4 # Provide the desired height for each small page inside the docx pages in INCHES


pdf_file_ext = '.pdf'
pdf_file_path = path+'/'+pdf_file_name+pdf_file_ext
 

if os.path.isdir(f"{path}/{pdf_file_name}_pictures"):

    """
    If there already exists a folder containing the images,
    we don't have to create the images from the PDF again.
    """

    docx_file_path = f"{path}/{pdf_file_name}_({w_i}x{h_i}).docx"
    shutil.copy2(standard_file_path, docx_file_path)

    total_pages = 0
    dir = f"{path}/{pdf_file_name}_pictures"
    for file_paths in os.listdir(dir):
        if os.path.isfile(os.path.join(dir, file_paths)):
            total_pages += 1

else:

    images = convert_from_path(pdf_file_path, poppler_path=r'C:/Program Files/poppler-0.68.0/bin') # convert a PDF file to individual images
    total_pages = len(images)
    print("Total Pages =", total_pages)

    os.mkdir(f"{path}/{pdf_file_name}_pictures") # create a folder to contain the images

    for i in range(len(images)):
        images[i].save(f'{path}/{pdf_file_name}_pictures/{pdf_file_name}({str(i+1)}).png', 'PNG') # save the images in that folder

    docx_file_path = f"{path}/{pdf_file_name}.docx" 
    shutil.copy2(standard_file_path, docx_file_path) # The provided docx file is copied and the images will be merged in that copy, leaving the given one untouched.


os.chdir(f"{path}/{pdf_file_name}_pictures")

document = Document(docx_file_path)



def pagemaker(pdf_file_name, i):

    """
    This is the main function which merges 8 pictures in 2 
    pages of the docx, in such a manner that 2 corresponding 
    pages of the PDF are printed exactly in one another's 
    opposite sides.
    """

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()

    if i==1:
        run.clear()

    try:
        run.add_picture(f'{pdf_file_name}({str(i)}).png', width= Inches(w_i), height= Inches(h_i))
    except:
        print(f'{pdf_file_name}({str(i)}).png not found')
    finally:
        i=i+2
        run.add_text('  ')

    try:
        run.add_picture(f'{pdf_file_name}({str(i)}).png', width= Inches(w_i), height= Inches(h_i))
    except:
        print(f'{pdf_file_name}({str(i)}).png not found')
    finally:
        i=i+2
        run.add_break()
        run.add_break()
        run.add_break()

    try:
        run.add_picture(f'{pdf_file_name}({str(i)}).png', width= Inches(w_i), height= Inches(h_i))
    except:
        print(f'{pdf_file_name}({str(i)}).png not found')
    finally:
        i=i+2
        run.add_text('  ')

    try:
        run.add_picture(f'{pdf_file_name}({str(i)}).png', width= Inches(w_i), height= Inches(h_i))
    except:
        print(f'{pdf_file_name}({str(i)}).png not found')
    finally:
        i=i-3

    document.add_page_break()

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()

    try:
        run.add_picture(f'{pdf_file_name}({str(i)}).png', width= Inches(w_i), height= Inches(h_i))
    except:
        print(f'{pdf_file_name}({str(i)}).png not found')
    finally:
        i=i-2
        run.add_text('  ')

    try:
        run.add_picture(f'{pdf_file_name}({str(i)}).png', width= Inches(w_i), height= Inches(h_i))
    except:
        print(f'{pdf_file_name}({str(i)}).png not found')
    finally:
        i=i+6
        run.add_break()
        run.add_break()
        run.add_break()

    try:
        run.add_picture(f'{pdf_file_name}({str(i)}).png', width= Inches(w_i), height= Inches(h_i))
    except:
        print(f'{pdf_file_name}({str(i)}).png not found')
    finally:
        i=i-2
        run.add_text('  ')

    try:
        run.add_picture(f'{pdf_file_name}({str(i)}).png', width= Inches(w_i), height= Inches(h_i))
    except:
        print(f'{pdf_file_name}({str(i)}).png not found')

    document.add_page_break()

    return


for i in range(1, total_pages+1, 8):
    """
    For every 8 pages, the main function is
    invoked once.
    """
    pagemaker(pdf_file_name, i)


document.save(docx_file_path)

print("Finished! program exiting now.")


